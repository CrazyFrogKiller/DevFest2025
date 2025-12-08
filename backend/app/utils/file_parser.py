from PyPDF2 import PdfReader
from docx import Document as DocxDocument
try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

class FileParser:
    """Utility class for parsing different file formats"""

    @staticmethod
    def parse_pdf(file_path: str) -> str:
        """Extract ALL text from PDF - maximum extraction mode"""
        try:
            text_parts = []

            # Метод 1: pdfplumber (лучший для сложных PDF)
            if HAS_PDFPLUMBER:
                try:
                    with pdfplumber.open(file_path) as pdf:
                        for page_num, page in enumerate(pdf.pages):
                            # Извлекаем весь текст с layout
                            page_text = page.extract_text(
                                layout=True,
                                x_tolerance=1,
                                y_tolerance=1
                            )

                            if page_text and page_text.strip():
                                text_parts.append(page_text.strip())

                            # Дополнительно извлекаем таблицы
                            tables = page.extract_tables()
                            if tables:
                                for table in tables:
                                    table_text = []
                                    for row in table:
                                        if row:
                                            row_text = ' | '.join([str(cell).strip() if cell else '' for cell in row])
                                            if row_text.strip():
                                                table_text.append(row_text)
                                    if table_text:
                                        text_parts.append('\n'.join(table_text))

                    if text_parts:
                        return "\n\n".join(text_parts)

                except Exception as e:
                    print(f"pdfplumber extraction failed: {e}")

            # Метод 2: PyPDF2 (fallback)
            text_parts = []
            with open(file_path, 'rb') as file:
                reader = PdfReader(file)

                for page_num, page in enumerate(reader.pages):
                    # Пробуем разные режимы извлечения
                    page_text = None

                    # Режим 1: layout (сохраняет форматирование)
                    try:
                        page_text = page.extract_text()
                    except:
                        pass

                    # Режим 2: plain (обычный)
                    if not page_text or not page_text.strip():
                        try:
                            page_text = page.extract_text()
                        except:
                            pass

                    if page_text and page_text.strip():
                        text_parts.append(page_text.strip())

            result = "\n\n".join(text_parts) if text_parts else ""

            if not result.strip():
                raise ValueError("PDF appears to be empty or contains only images")

            return result

        except Exception as e:
            raise ValueError(f"Failed to parse PDF: {str(e)}")

    @staticmethod
    def parse_docx(file_path: str) -> str:
        """Extract ALL text from DOCX - maximum extraction mode"""
        try:
            doc = DocxDocument(file_path)
            text_parts = []

            # 1. Заголовки документа (headers)
            for section in doc.sections:
                if section.header:
                    for paragraph in section.header.paragraphs:
                        text = paragraph.text.strip()
                        if text:
                            text_parts.append(f"[HEADER] {text}")

            # 2. Основной текст (paragraphs)
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_parts.append(text)

            # 3. Таблицы (tables) - КРИТИЧНО!
            for table in doc.tables:
                table_texts = []
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        # Извлекаем текст из каждой ячейки
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_texts.append(cell_text)

                    if row_texts:
                        table_texts.append(' | '.join(row_texts))

                if table_texts:
                    text_parts.append('\n'.join(table_texts))

            # 4. Нижние колонтитулы (footers)
            for section in doc.sections:
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        text = paragraph.text.strip()
                        if text:
                            text_parts.append(f"[FOOTER] {text}")

            # 5. Текстовые блоки (text frames) - если есть
            try:
                for shape in doc.inline_shapes:
                    if hasattr(shape, 'text_frame'):
                        for paragraph in shape.text_frame.paragraphs:
                            text = paragraph.text.strip()
                            if text:
                                text_parts.append(text)
            except:
                pass  # Не все документы имеют shapes

            result = "\n\n".join(text_parts)

            if not result.strip():
                raise ValueError("DOCX appears to be empty")

            return result

        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {str(e)}")

    @staticmethod
    def parse_txt(file_path: str) -> str:
        """Read text from TXT file"""
        try:
            encodings = ['utf-8', 'utf-8-sig', 'cp1251', 'windows-1251', 'latin-1', 'cp1252']

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except (UnicodeDecodeError, LookupError):
                    continue

            # Последняя попытка
            with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                return file.read()

        except Exception as e:
            raise ValueError(f"Failed to parse TXT: {str(e)}")

    @staticmethod
    def parse_md(file_path: str) -> str:
        """Read markdown file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            raise ValueError(f"Failed to parse Markdown: {str(e)}")

    @staticmethod
    def parse_file(file_path: str, file_type: str) -> str:
        """Parse file based on its type"""
        file_type = file_type.lower()

        if file_type == 'pdf':
            return FileParser.parse_pdf(file_path)
        elif file_type == 'docx':
            return FileParser.parse_docx(file_path)
        elif file_type == 'txt':
            return FileParser.parse_txt(file_path)
        elif file_type == 'md':
            return FileParser.parse_md(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")